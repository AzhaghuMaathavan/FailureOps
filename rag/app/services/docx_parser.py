import uuid
import logging
from app.models.document import Page, DocumentBlock

logger = logging.getLogger(__name__)

def parse_docx_to_blocks(file_path: str, doc_id: str, db):
    try:
        import docx
        doc = docx.Document(file_path)
    except ImportError:
        logger.warning("python-docx not installed. Skipping docx parsing.")
        return False
    except Exception as e:
        logger.error(f"Failed to load DOCX {file_path}: {e}")
        return False

    page_id = str(uuid.uuid4())
    db_page = Page(
        id=page_id,
        document_id=doc_id,
        page_number=1,
        image_path="",
        status="COMPLETED"
    )
    db.add(db_page)

    block_idx = 0
    current_section = "General"
    
    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue
            
        b_type = "Paragraph"
        if p.style and p.style.name and p.style.name.startswith("Heading"):
            b_type = "Section-header"
            current_section = text
            
        db_block = DocumentBlock(
            id=str(uuid.uuid4()),
            document_id=doc_id,
            page_id=page_id,
            block_index=block_idx,
            block_type=b_type,
            content=text,
            raw_metadata={"section": [current_section]}
        )
        db.add(db_block)
        block_idx += 1

    for t in doc.tables:
        lines = []
        for r_idx, row in enumerate(t.rows):
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            lines.append(f"| {' | '.join(cells)} |")
            if r_idx == 0:
                lines.append(f"| {' | '.join(['---'] * len(cells))} |")
                
        table_md = "\n".join(lines)
        db_block = DocumentBlock(
            id=str(uuid.uuid4()),
            document_id=doc_id,
            page_id=page_id,
            block_index=block_idx,
            block_type="Table",
            content=table_md,
            raw_metadata={"section": [current_section]}
        )
        db.add(db_block)
        block_idx += 1

    db.commit()
    return True
